import docx
import plotly as plt
import plotly.express as px
import json
from urllib.request import urlopen
import pandas as pd
import io
import matplotlib.pyplot as plt
import shapefile as shp
#matplotlib inline
import seaborn as sns
import geopandas as gpd
import numpy as np

doc = docx.Document('tabela2.docx')

try:
    document = docx.Document('tabela2.docx')
except:
    document = docx.Document()
    document.save('tabela2.docx')
    print("Previous file was corrupted or didn't exist - new file was created.")

count = 0 
for table in doc.tables:
    print(table)
    for row in table.rows:
        count += 1
    for cell in row.cells:
        for para in cell.paragraphs:
            print(para.text)
            print(count)

documento = docx.Document('tabela2.docx')

table = documento.tables[0]
print(table)

data = []
keys = None
for i, row in enumerate(table.rows):
    text = (cell.text for cell in row.cells)

keys = None
for i, row in enumerate(table.rows):
    text = (cell.text for cell in row.cells)
    print(i, row)

    if i == 0:
        keys = tuple(text)
        continue
    row_data = dict(zip(keys, text))
    if row_data not in data:
        data.append(row_data)
    else:
        break

tbl_01 = pd.DataFrame(data=data)
tbl_01
print(tbl_01)

tb = []
for t in tbl_01.columns.values:
    if t != '' and len(t) < 4:
        a = int(t)
        a.append(tb)

total_uber = 0 
for k, v in tbl_01['UBER'].items():
    if v != '':
        total_uber += int(v)

total_99 = 0 
for k, v in tbl_01.iloc[:, 2 ].items():
    if v != '':
        total_99 += int(v)

total_ifood = 0 
for k, v in tbl_01['IFOOD'].items():
    if v != '':
        total_ifood += int(v)

total_loggi = 0 
for k, v in tbl_01['LOGGI'].items():
    if v != '':
        total_loggi += int(v)

total_rappi = 0 
for k, v in tbl_01['RAPPI'].items():
    if v != '':
        total_rappi += int(v)

tbl_01.loc[len(tbl_01.index)] = ['Total', total_uber, total_99, total_ifood, total_loggi, total_rappi]

print(tbl_01)

#GRÁFICO 

# libraries
#import numpy as np
#import matplotlib.pyplot as plt
#from matplotlib import rc
#import pandas as pd

# Data
#r = [0,1,2,3,4]
#SCLICING - COMPANIES - ROWS - LINHAS
#uber = tbl_01.UBER
#taxis_99 = tbl_01['99 TAXIS']
#ifood = tbl_01['IFOOD']
#loggi = tbl_01['LOGGI']
#rappi = tbl_01['RAPPI']



#REGIOES DO BRASIL 
#norte = tbl_01[7, 9, 10, 13]
#nordeste = tbl_01[4, 5, 6, 12, 15, 18, 19, 20, 21]
#centroeste = tbl_01[17, 22, 23]
#sudeste = tbl_01[0, 1, 2, 14, 16]
#sul = tbl_01[3, 8, 11]


#raw_data = {'greenBars': uber, 'orangeBars': taxis_99,'blueBars':ifood}
#df = pd.DataFrame(raw_data)

# From raw value to percentage
#totals = [i+j+k for i,j,k in zip(df['greenBars'], df['orangeBars'], df['blueBars'])]
#greenBars = [i / j * 100 for i,j in zip(df['greenBars'], totals)]
#orangeBars = [i / j * 100 for i,j in zip(df['orangeBars'], totals)]
#blueBars = [i / j * 100 for i,j in zip(df['blueBars'], totals)]

# plot
#barWidth = 0.85
#names = ('A','B','C','D','E')
# Create green Bars
#plt.bar(r, greenBars, color='#b5ffb9', edgecolor='white', width=barWidth)
# Create orange Bars
#plt.bar(r, orangeBars, bottom=greenBars, color='#f9bc86', edgecolor='white', width=barWidth)
# Create blue Bars
#plt.bar(r, blueBars, bottom=[i+j for i,j in zip(greenBars, orangeBars)], color='#a3acff', edgecolor='white', width=barWidth)

# Custom x axis
#plt.xticks(r, names)
#plt.xlabel("group")

# Show graphic
#plt.show()
#INFOS_UFS = gpd.read_file('bcim_2016_21_11_2018.gpkg', layer = 'lim_unidade_federacao_a')

#Colunas existentes no dataframe
#INFOS_UFS.columns

#print("Estado: " + INFOS_UFS.nome[2])
#INFOS_UFS.geometry[2]

#ADEQUAR OS TRTS PARA OS ESTADOS CORRESPONDENTES

trt_estadoc = {'TRT1': ['RJ'], 'TRT2': ['SP'], 'TRT3': ['MG'], 'TRT4': ['RS'], 'TRT5': ['BA'], 'TRT6': ['PE'], 'TRT7': ['CE'], 'TRT8': ['AP', 'PA'], 'TRT9': ['PR'], 'TRT10':['TO', 'DF'], 'TRT11':['AM', 'RR'], 'TRT12':['SC'], 'TRT13': ['PB'], 'TRT14':['AC', 'RO'],'TRT15':['SP'], 'TRT16':['MA'],'TRT17': ['ES'], 'TRT18': ['GO'], 'TRT19': ['AL'], 'TRT20': ['SE'], 'TRT21':['RN'], 'TRT22':['PI'], 'TRT23': ['MT'], 'TRT24':['MS'], 'TST': ['']}
#####trt_estadoc = {'RJ': ['TRT1'], 'SP': ['TRT2', 'TRT15'], 'MG': ['TRT3'], 'RS':['TRT4'], 'BA': ['TRT5'], 'PE': ['TRT6'], 'CE': ['TRT7'], 'AP': ['TRT8'], 'PA': ['TRT8'], 'PR': ['TRT9'], 'TO':['TRT10'], 'DF':['TRT10'], 'AM':['TRT11'], 'RR':['TRT11'], 'SC':['TRT12'], 'PB': ['TRT13'], 'AC':['TRT14'], 'RO':['TRT14'], 'MA': ['TRT16'], 'ES':['TRT17'], 'GO':['TRT18'], 'AL':['TRT19'], 'SE':['TRT20'], 'RN':['TRT21'], 'PI': ['TRT22'], 'MT': ['TRT23'], 'MS':['TRT24'], 'BR': [''] }
#estadoc = '' 
#for k, v in trt_estadoc.items():
#    b = ', '.join(v)
#    estadoc += b 

#my_string = ', '.join(map(str, trt_estadoc.values()))

#===============================
#SOMA DAS LINHAS - TRIBUNAIS E PROCESSOS TOTAIS - CEAT ==============


intervo = tbl_01.loc[1:,"UBER":"RAPPI"]
#tbl_01['Resultados - Soma'] = tbl_01[intervo].sum(axis=1)
#print(tbl_01)



print(tbl_01)







#=========================================
tbl_01['sigla'] = trt_estadoc.values()

todose = []
for estado in tbl_01['sigla']:
    eststring = ' '
    ea = str(estado)
    if ea not in eststring:
        eststring += ea
    todose.append(eststring)
    #print(type(estado))
    #ea = str(estado)
    #print(type(ea))
    #print(ea)
#print(todose)

#for item in todose:
    #print(type(item))

novo = []
for x in todose:
    item = x
    for y in ['\n', '\t', '/', '.', '-', '(', ')', '[', ']', "'"]:
        item = item.replace(y, "")
    novo.append(item)


tbl_01.insert(1, 'state', novo, True)

tbl_01.drop('sigla', inplace=True, axis=1)

print(tbl_01)

#INFOS_UFS.rename({'sigla': 'state'}, axis = 1, inplace = True)

#BRASIL = INFOS_UFS.merge(tbl_01, on = 'state', how = 'left')

#print(BRASIL[['state', 'UBER', '99 TAXIS', 'IFOOD', 'LOGGI', 'RAPPI', 'geometry' ]].head())

#BRASIL.plot(column = 'UBER', cmap = 'Reds', figsize = (16,10), legend = True, edgecolor = 'black')

# TENTATIVA DE GRÁFICO COM PLOTY.EX - BAR PLOTS 
#import plotly.express as px

#fig = px.bar(tbl_01, x='UBER', y= '', orientation='h')
#fig.show()

#TENTATIVA 2




#df2 = pd.DataFrame([uber], columns=[uber])

#df2.plot.bar()

df_tr =tbl_01.transpose()

print(df_tr)


df_trans = df_tr.iloc[2:, 0:25]

print(df_trans)

total_trts = []
soma_por_linha = []
for k, v in df_trans.items():
    soma_linha = 0 
    if 'TRT' not in v:
            for o  in v:
                if o != '' :
                    a = int(o)
                    soma_linha += a 
                    total_trts.append(a) 
            soma_por_linha.append(soma_linha)    


soma_total_trts = sum(total_trts)

#print("TABLEA COM EMPRESAS/totais:", tbl_01)
print("Essa é a lista dos trts: ", total_trts, "e este o total por linha: ", soma_por_linha, "essa é a soma de tudo:", soma_total_trts )


df_trans.loc['TRTs Total'] = soma_por_linha

print(df_trans)


tbl_01['TOTAL P/ TRT'] = soma_por_linha

print(tbl_01)

df_trans.to_csv('tabela_02_processos.csv')
writer = pd.ExcelWriter('tabela_02_processos.xlsx')
df_trans.to_excel(writer)
# save the excel
writer.save()
print('DataFrame is written successfully to Excel File.')

tbl_01.to_csv('tabela_01_processos.csv')
writer = pd.ExcelWriter('tabela_01_processos.xlsx')
tbl_01.to_excel(writer)
writer.save()
print('DataFrame is written successfully to Excel File.')

import plotly.express as px
df = tbl_01.iloc[:24, :]
#interativo = ['UBER', '99 TAXIS', 'IFOOD', 'LOGGI', 'RAPPI']
#fig = px.bar(df, x='TOTAL P/ TRT', y='99 TAXIS', color='', orientation = 'h',
            #hover_data=["TOTAL", "size"],
            #height=800,
            #title='CEATS Empresas - Tribunais')
#fig.show()

fig = px.bar(df, x='', y='TOTAL P/ TRT',
            #hover_data=['UBER', 'TAXIS99', 'IFOOD', 'LOGGI', 'RAPPI'], 
            color='TOTAL P/ TRT',
            #labels={'pop':'population of Canada'}, 
            height=400)
fig.show()


import os

from kaleido.scopes.plotly import PlotlyScope
import plotly.graph_objects as go
scope = PlotlyScope(
    plotlyjs="https://cdn.plot.ly/plotly-latest.min.js",
    # plotlyjs="/path/to/local/plotly.js",
)

#==== SALVAR IMAGEM COM ESTE CÓDIGO - QUANDO FOR SALVAR, TIRAR O COMENTÁRIO DO CÓDIGO
with open("figure.jpg", "wb") as f:
    f.write(scope.transform(fig, format="jpg"))


fig.write_html("grafico_01.html")

import plotly.graph_objects as go
from plotly.subplots import make_subplots



# Create subplots: use 'domain' type for Pie subplot
fig = make_subplots(rows=1, cols=2, specs=[[{'type':'domain'}, {'type':'domain'}]])
fig.add_trace(go.Pie(labels=df[''], values=df['UBER'], name="UBER"),
            1, 1)
fig.add_trace(go.Pie(labels=df[''], values=df.iloc[:, 3], name="TAXIS99"),
            1, 2)

# Use `hole` to create a donut-like pie chart
fig.update_traces(hole=.4, hoverinfo="label+percent+name")

fig.update_layout(
    title_text="Porcentagens - Processos",
    # Add annotations in the center of the donut pies.
    annotations=[dict(text='UBER', x= 0.2, y=0.5, font_size=20, showarrow=False),
                dict(text='TAXIS99', x=0.81, y=0.5, font_size=20, showarrow=False)])
fig.show()


fig.write_html("grafico_02.html")


#fig.show()


#df['state'].value_counts()


#GŔAFICO DE COLUNAS COM GRUPOS 

import plotly.graph_objects as go

trts = ['TRT1', 'TRT2', 'TRT3', 'TRT4', 'TRT5', 'TRT6',
        'TRT7', 'TRT8', 'TRT9', 'TRT10', 'TRT11', 'TRT12', 
        'TRT13', 'TRT14', 'TRT15', 'TRT16', 'TRT17', 'TRT18', 
        'TRT19', 'TRT20', 'TRT21', 'TRT22', 'TRT23', 'TRT24', 'TOTAL']

fig = go.Figure()
fig.add_trace(go.Bar(
    x=trts,
    y=tbl_01['UBER'],
    name='Uber',
    marker_color='indianred'
))
fig.add_trace(go.Bar(
    x=trts,
    y=tbl_01.iloc[:, 3],
    name='Taxis99',
    marker_color='lightsalmon'
    
))

# Here we modify the tickangle of the xaxis, resulting in rotated labels.
fig.update_layout(barmode='group', xaxis_tickangle=-45)
fig.show()
fig.write_html("grafico_transporte.html")

with open("processos_uber-99.jpg", "wb") as f:
    f.write(scope.transform(fig, format="jpg"))

print(df_trans)

import matplotlib.pyplot as plt


data=tbl_01

df=pd.DataFrame(data,index=trts)

df.plot(kind="bar",stacked=True)
plt.legend(loc="lower left",bbox_to_anchor=(0.8,1.0))
plt.show()





