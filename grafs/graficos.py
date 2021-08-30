import docx
import plotly as plt
import plotly.express as px
import json
from urllib.request import urlopen
import pandas as pd
import io
import matplotlib.pyplot as plt
import shapefile as shp
#matplotlib inline
import seaborn as sns
import geopandas as gpd
import numpy as np

doc = docx.Document('tabela2.docx')

try:
    document = docx.Document('tabela2.docx')
except:
    document = docx.Document()
    document.save('tabela2.docx')
    print("Previous file was corrupted or didn't exist - new file was created.")

count = 0 
for table in doc.tables:
    print(table)
    for row in table.rows:
        count += 1
    for cell in row.cells:
        for para in cell.paragraphs:
            print(para.text)
            print(count)

documento = docx.Document('tabela2.docx')

table = documento.tables[0]
print(table)

data = []
keys = None
for i, row in enumerate(table.rows):
    text = (cell.text for cell in row.cells)

keys = None
for i, row in enumerate(table.rows):
    text = (cell.text for cell in row.cells)
    print(i, row)

    if i == 0:
        keys = tuple(text)
        continue
    row_data = dict(zip(keys, text))
    if row_data not in data:
        data.append(row_data)
    else:
        break

tbl_01 = pd.DataFrame(data=data)
tbl_01
print(tbl_01)

#==============================================================================

import plotly.express as px
import os

from kaleido.scopes.plotly import PlotlyScope
import plotly.graph_objects as go
scope = PlotlyScope(
    plotlyjs="https://cdn.plot.ly/plotly-latest.min.js",
    # plotlyjs="/path/to/local/plotly.js",
)

fig = px.bar(tbl_01, x="TRIBUNAIS", y="PROCESSOS", color="EMPRESAS", title="Pesquisa de Processos Através de CEATs")
fig.show()

fig.write_html("grafico_ceat_stacked-bar.html")

with open("grafico_ceat_stacked-bar.jpg", "wb") as f:
    f.write(scope.transform(fig, format="jpg"))
